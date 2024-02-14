from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def editFormat(DATE,TIME):
    template_file_path = 'Format v2\\ticket version 71.docx'    

    variables = {
        "$DATE": DATE,
        "$TIME": TIME,                
    }

    template_document = Document(template_file_path)
    # for variable_key, variable_value in variables.items():
    #     for paragraph in template_document.paragraphs:
    #         replace_text_in_paragraph(paragraph, variable_key, variable_value)
    

    old_text = "DATE"
    new_text = DATE

    # Iterate through paragraphs and search for the text
    # for paragraph in template_document.paragraphs:
    #     if old_text in paragraph.text:
    #         for run in paragraph.runs:
    #             if old_text in run.text:
    #                 run.text = run.text.replace(old_text, new_text)

    # Iterate through text boxes and shapes
    for shape in template_document.pict:
        if hasattr(shape, "text"):
            if old_text in shape.text:
                shape.text = shape.text.replace(old_text, new_text)


    # Function to recursively search and replace text within style boxes
    # def find_and_replace_in_style_boxes(container):
    #     for paragraph in container.paragraphs:
    #         for run in paragraph.runs:
    #             run.text = run.text.replace(old_text, new_text)
        
    #     for table in container.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 find_and_replace_in_style_boxes(cell)

    # find_and_replace_in_style_boxes(template_document)

    template_document.save(f'Edited Docx v2\\Edited.docx')
    print('\n\t\tPoster Created!')    

# def replace_text_in_paragraph(paragraph, key, value):
#     if key in paragraph.text:
#         inline = paragraph.runs
#         for item in inline:
#             if key in item.text:
#                 item.text = item.text.replace(key, value)


if __name__ == '__main__':
    editFormat('13.14.2012','12:34:45')